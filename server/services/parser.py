import re
import io
import time
from typing import List, Optional, Tuple
from pathlib import Path

import pdfplumber
from docx import Document

from server.models import Resume, SkillItem, Experience
from server.services.skill_dict import skill_dict, EDUCATION_LEVELS


PHONE_PATTERN = re.compile(r"(?<!\d)(1[3-9]\d{9})(?!\d)")
EMAIL_PATTERN = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
YEARS_PATTERN = re.compile(r"(\d+(?:\.\d+)?)\s*(?:年|years?|yrs?)", re.IGNORECASE)
EDUCATION_PATTERN = re.compile(
    r"(博士|硕士|研究生|本科|学士|大专|专科|高中|PhD|Doctorate|Master|Bachelor|Associate)",
    re.IGNORECASE,
)
CHINESE_NAME_PATTERN = re.compile(r"^[\u4e00-\u9fa5]{2,4}$")


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    text_parts: List[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            text_parts.append(t)
    return "\n".join(text_parts).strip()


def _extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text.strip())
    return "\n".join(lines).strip()


def _extract_text_from_txt(file_bytes: bytes) -> str:
    for enc in ("utf-8", "gbk", "gb18030", "utf-16"):
        try:
            return file_bytes.decode(enc).strip()
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="ignore").strip()


def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _extract_text_from_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return _extract_text_from_docx(file_bytes)
    elif ext in (".txt", ".md", ".text"):
        return _extract_text_from_txt(file_bytes)
    else:
        raise ValueError(f"不支持的文件格式: {ext}")


def _guess_name(text: str, filename: str) -> Optional[str]:
    lines = [l.strip() for l in text.splitlines()[:30] if l.strip()]
    for line in lines:
        tokens = re.split(r"[\s,，·|/\\]+", line)
        for tok in tokens:
            if CHINESE_NAME_PATTERN.match(tok):
                return tok
    stem = Path(filename).stem
    if CHINESE_NAME_PATTERN.match(stem):
        return stem
    if lines:
        first = lines[0]
        if 2 <= len(first) <= 20 and not PHONE_PATTERN.search(first) and not EMAIL_PATTERN.search(first):
            return first[:20]
    return None


def _extract_phone(text: str) -> Optional[str]:
    m = PHONE_PATTERN.search(text)
    return m.group(1) if m else None


def _extract_email(text: str) -> Optional[str]:
    m = EMAIL_PATTERN.search(text)
    return m.group(0).lower() if m else None


def _extract_education(text: str) -> Optional[str]:
    found = EDUCATION_PATTERN.findall(text)
    if not found:
        return None
    best = None
    best_level = 0
    for f in found:
        lower = f.lower()
        for key, lv in EDUCATION_LEVELS.items():
            if key.lower() in lower or lower in key.lower():
                if lv > best_level:
                    best_level = lv
                    best = key
                    break
    return best


def _extract_years(text: str) -> int:
    nums = []
    for m in YEARS_PATTERN.finditer(text):
        try:
            nums.append(float(m.group(1)))
        except ValueError:
            pass
    if not nums:
        return 0
    nums.sort()
    probable = [n for n in nums if 1 <= n <= 40]
    if not probable:
        return 0
    return int(round(max(probable)))


def _extract_skills(text: str) -> Tuple[List[SkillItem], List[str]]:
    extracted: List[SkillItem] = []
    seen_canonical = set()
    candidates_new = []
    tokens = re.split(r"[\s,，、;；/|\\()（）\[\]【】]+", text)
    for tok in tokens:
        t = tok.strip()
        if len(t) < 2 or len(t) > 30:
            continue
        canon = skill_dict.normalize(t)
        if canon and canon not in seen_canonical:
            extracted.append(SkillItem(name=canon, confidence=1.0))
            seen_canonical.add(canon)
        elif not canon and re.search(r"[A-Za-z\u4e00-\u9fa5]", t):
            if re.search(r"[A-Za-z]", t) or (len(t) <= 8 and re.search(r"[\u4e00-\u9fa5]", t)):
                candidates_new.append(t)
    skill_section = re.search(
        r"(?:技能|技术|专业技能|掌握|熟悉|技术栈|skill|skills|technology|tech stack)[\s:：]*([\s\S]{0,800})",
        text, re.IGNORECASE,
    )
    if skill_section:
        section_text = skill_section.group(1)
        for tok in re.split(r"[\s,，、;；/|\\()（）\[\]【】]+", section_text):
            t = tok.strip()
            if len(t) < 2:
                continue
            canon = skill_dict.normalize(t)
            if canon and canon not in seen_canonical:
                extracted.append(SkillItem(name=canon, confidence=1.0))
                seen_canonical.add(canon)
    return extracted, list(set(candidates_new))


def _extract_companies(text: str) -> List[Experience]:
    results: List[Experience] = []
    exp_section = re.search(
        r"(?:工作经历|工作经验|工作|职业经历|professional experience|work experience|experience)[\s:：]*([\s\S]{0,2000})",
        text, re.IGNORECASE,
    )
    source = exp_section.group(1) if exp_section else text
    lines = [l.strip() for l in source.splitlines() if l.strip()]
    for line in lines[:30]:
        if len(line) > 100:
            continue
        company_matches = re.findall(
            r"([\u4e00-\u9fa5A-Za-z][\u4e00-\u9fa5A-Za-z0-9·\-\s（）()]{2,40}?(?:公司|科技|集团|有限公司|股份|Inc|Corp|Ltd|LLC|Bank|Lab|实验室|大学|学院|School|University))",
            line,
        )
        for c in company_matches:
            c = c.strip()
            if any(not e.company or e.company != c for e in results):
                if not any(e.company == c for e in results):
                    results.append(Experience(company=c))
            if len(results) >= 3:
                return results
    return results[:3]


def parse_resume(file_bytes: bytes, filename: str) -> Resume:
    start = time.time()
    errors: List[str] = []
    try:
        text = extract_text(file_bytes, filename)
    except Exception as e:
        raise ValueError(f"文件解析失败: {e}")
    if not text or len(text.strip()) < 20:
        raise ValueError("无法识别文字内容或内容过短")
    text = text.strip()
    name = _guess_name(text, filename)
    phone = _extract_phone(text)
    email = _extract_email(text)
    education = _extract_education(text)
    years = _extract_years(text)
    skills, new_candidates = _extract_skills(text)
    companies = _extract_companies(text)

    confidence = 0.0
    checks = 0
    if name:
        confidence += 0.2
        checks += 1
    if phone:
        confidence += 0.2
        checks += 1
    if email:
        confidence += 0.2
        checks += 1
    if education:
        confidence += 0.15
        checks += 1
    if years > 0:
        confidence += 0.1
        checks += 1
    if skills:
        confidence += 0.1 * min(len(skills), 15) / 15
        checks += 1
    if companies:
        confidence += 0.05
        checks += 1
    confidence = min(round(confidence, 3), 1.0)

    resume = Resume(
        filename=filename,
        name=name,
        phone=phone,
        email=email,
        education=education,
        years_of_experience=years,
        skills=skills,
        recent_companies=companies,
        raw_text=text,
        confidence=confidence,
    )

    for term in new_candidates:
        skill_dict.record_candidate(term)

    return resume
